import os
import sqlite3
import fitz  # PyMuPDF
from docx import Document
import re

# --- 1. 路徑設定 (確保與 main.py 同步) ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# 法規檔案放在同目錄下的 data 資料夾
SOURCE_DIR = os.path.join(BASE_DIR, "data")
# 資料庫檔案路徑
DB_PATH = os.path.join(BASE_DIR, "legal_index.db")

def get_category(filename):
    """根據檔名自動分類"""
    categories = {
        "法": "法律",
        "辦法": "辦法",
        "要點": "要點",
        "規則": "規則",
        "章程": "章程",
        "簡則": "簡則",
        "須知": "作業須知",
        "手冊": "手冊",
        "範例": "範例"
    }
    for kw, cat in categories.items():
        if kw in filename:
            return cat
    return "其他"

def extract_text_pdf(path):
    """提取 PDF 文字內容"""
    try:
        doc = fitz.open(path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        print(f"❌ PDF 解析失敗 {os.path.basename(path)}: {e}")
        return ""

def extract_text_docx(path):
    """提取 Word 文字內容 (包含段落與表格)"""
    try:
        doc = Document(path)
        full_text = []
        # 1. 讀取所有段落
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 2. 讀取所有表格 (法規中常有表格內容，這部分很重要)
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    full_text.append(" | ".join(row_data))
                    
        return "\n".join(full_text)
    except Exception as e:
        print(f"❌ Word 解析失敗 {os.path.basename(path)}: {e}")
        return ""

def init_db():
    """初始化 SQLite FTS5 全文檢索資料庫"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    # 每次重建索引，確保資料最新且沒有重複
    cursor.execute("DROP TABLE IF EXISTS docs")
    cursor.execute("""
        CREATE VIRTUAL TABLE docs USING fts5(
            title,
            category,
            content,
            path,
            filename UNINDEXED
        )
    """)
    conn.commit()
    return conn

def run_indexer():
    """執行索引更新主程式"""
    print("\n🚀 --- 開始建立法規索引 ---")
    
    # 檢查資料夾是否存在
    if not os.path.exists(SOURCE_DIR):
        print(f"📁 建立資料夾: {SOURCE_DIR}")
        os.makedirs(SOURCE_DIR, exist_ok=True)
    
    conn = init_db()
    cursor = conn.cursor()
    
    # 取得檔案清單
    try:
        files = [f for f in os.listdir(SOURCE_DIR) if f.lower().endswith(('.pdf', '.docx'))]
    except Exception as e:
        print(f"❌ 無法讀取資料夾: {e}")
        files = []
        
    print(f"📂 找到 {len(files)} 份法規檔案於 {SOURCE_DIR}")
    
    indexed_count = 0
    for filename in files:
        path = os.path.join(SOURCE_DIR, filename)
        category = get_category(filename)
        title = os.path.splitext(filename)[0]
        
        print(f"⏳ [{indexed_count + 1}/{len(files)}] 正在處理: {filename}")
        
        try:
            if filename.lower().endswith('.pdf'):
                content = extract_text_pdf(path)
            else:
                content = extract_text_docx(path)
                
            if content and content.strip():
                cursor.execute(
                    "INSERT INTO docs (title, category, content, path, filename) VALUES (?, ?, ?, ?, ?)",
                    (title, category, content, path, filename)
                )
                indexed_count += 1
            else:
                print(f"⚠️  警告: {filename} 內容為空，跳過索引。")
                
        except Exception as e:
            print(f"❌ 略過 {filename}，錯誤訊息: {e}")
        
    conn.commit()
    conn.close()
    print(f"✅ --- 索引建立完成！成功處理 {indexed_count} 份法規 ---\n")

if __name__ == "__main__":
    run_indexer()
